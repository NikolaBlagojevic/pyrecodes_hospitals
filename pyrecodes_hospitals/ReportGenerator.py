import docx
import os
from pyrecodes_hospitals import Plotter
from pyrecodes_hospitals import ResilienceCalculator
from pyrecodes_hospitals import main
import matplotlib.pyplot as plt
import matplotlib.ticker
import MCI_Planning_Tool_GUI

MEASURES_SERVICE_IN_TABLES = ['MortalityRateBefore24H', 'MortalityRateAfter24H', 'AverageLengthOfStay', 'SurgeriesPerformed', 'SurgeriesCancelled']
MEASURES_SERVICE_IN_TABLES_LABELS = ['Mortality Rate Before 24 Hours', 'Mortality Rate After 24 Hours', 'Average Length of Stay', 'Surgeries Performed', 'Surgeries Cancelled']
MEASURES_SERVICE_IN_TABLES_UNITS = ['%', '%', 'hours']
RESOURCE_NAME_IN_EXCEL = ['Nurses', 'Power', 'Water', 'Oxygen', 'Medical drugs', 'Blood', 'Beds']
MEASURES_SERVICE_NOTES = {'MortalityRateBefore24H': 'The mortality rate before 24 hours is the percentage of patients that died within the first 24 hours of admission divided by the total number of patients of the considered patient type.',
                          'MortalityRateAfter24H': 'The mortality rate after 24 hours is the percentage of patients that died after 24 hours of admission divided by the total number of patients of the considered patient type.',
                          'AverageLengthOfStay': 'The average length of stay is the average number of hours a patient of the considered patient type stays in the hospital. Please note that the length of stay can depend on the considered duration of the MCI simulation as defined in the stress scenario (e.g., if the considered duration is 2 days, no length of stay can be longer than 2 days).'}
MODIFIED_RESOURCE_NAME = {'Nurse': 'Nurses',
                          'Fuel': 'Fuel',
                          'Water': 'Water',
                          'Oxygen': 'Oxygen',
                          'MedicalDrugs': 'Medical drugs',
                          'Blood': 'Blood',
                          'Stretcher': 'Stretchers',
                          'MCI_Kit_Walking_RestOfHospital': 'MCI kits to treat walking patients in the Emergency Department',
                          'MCI_Kit_NonWalking_EmergencyDepartment': 'MCI kits to treat non-walking patients in the Emergency Department',
                          'MCI_Kit_NonWalking_OperatingTheater': 'MCI kits to treat patients in the Operating Theater',
                          'MCI_Kit_NonWalking_HighDependencyUnit': 'MCI kits to treat patients in the High Dependency Unit',
                          'MCI_Kit_NonWalking_Medical/SurgicalDepartment': 'MCI kits to treat patients in the Medical/Surgical Department',
                          'EmergencyDepartment_Bed': 'equipped and staffed beds in the Emergency Department',
                          'OperatingTheater_Bed': 'equipped and staffed beds in the Operating Theater',
                          'Medical/SurgicalDepartment_Bed': 'equipped and staffed beds in the Medical/Surgical Department',
                          'HighDependencyUnit_Bed': 'equipped and staffed beds in the High Dependency Unit',
                          'RestOfHospital_Bed': 'equipped and staffed beds in rest of the hospital'
                          }

CONSUMABLE_RESOURCES = ['MCI_Kit_Walking_RestOfHospital', 'MCI_Kit_NonWalking_EmergencyDepartment',
                        'MCI_Kit_NonWalking_OperatingTheater', 'MCI_Kit_NonWalking_HighDependencyUnit',
                        'MCI_Kit_NonWalking_Medical/SurgicalDepartment', 'Blood']

def generate_report(system, input_file_location: str):
    doc = docx.Document()

    # Add a title
    doc.add_heading('Hospital Resilience Assessment', 0)
    doc.add_heading('MCI Planning Tool', level=1)

    # Add description of the MCI Planning Tool
    doc.add_paragraph('This report is generated using the MCI Planning Tool. MCI Planning Tool simulates the flow of patients through hospital departments during a mass casualty incident (MCI). It estimates the demand for resources at a department, compares it with the available supply and estimates the impact on patients in cases where supply is not sufficient to meet the demand. 7 resource categories are considered in this version of the tool: Fuel, Water, Nurses, Oxygen, Beds and Stretchers, Blood and Medical Drugs/Kits. More comprehensive hospital performance assessment is possible using the Rapid and Detailed Resilience Assessment Tools. The MCI Planning Tool is built on several conservative assumptions regarding the hospital performance in an MCI, not considering substandard procedures. The outputs of the tool should thus be used as a conservative and approximate estimate of the hospital performance and need to be assessed by experienced hospital staff and decision makers prior to implementing interventions in the hospital based on the tool outputs.')

    doc.add_heading('Results of MCI resilience assessment', level=1)

    doc.add_paragraph('In the following section, the results of the MCI resilience assessment are summarized. The results are presented as measures of service for each department and patient type in the hospital. The measures of service include the mortality rate before and after 24 hours, surgeries performed and cancelled, and the average length of stay. In addition to the measures of service, causes of death are identified by analyzing the unmet resource demand of patients that died. Finally, to understand resource bottlenecks, the supply and demand of resources over time during the MCI are visualized in the form of plots.')

    add_measures_of_service_summary(system, doc)

    add_cause_of_death_summary(system, doc)

    add_supply_demand_plots_to_report(system, doc)

    add_liability_waiver(doc)

    doc.add_heading('Annex: Inputs', level=1)

    doc.add_paragraph('The following section presents the input values used to run the MCI resilience assessment in the MCI Planning Tool. The input values include the resource supply, patient profiles, and the stress scenario.')

    add_resource_supply_description(input_file_location, doc)

    add_stress_scenario_description(system, doc)

    add_patient_profiles_description(system, doc)

    add_contact_information(doc)

    # Justify all paragraphs
    for paragraph in doc.paragraphs:
        paragraph.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # Save the document
    doc.save('MCIResilienceAssessmentReport.docx')

    if os.path.exists('MCIResilienceAssessmentReport.docx'):
        return os.path.abspath('MCIResilienceAssessmentReport.docx')
    else:
        return None

def add_resource_supply_description(input_file_location, doc):

    doc.add_heading('Hospital Resource Supply', level=2)

    doc.add_paragraph('The following tables present the resource supply values used as the input for the MCI Planning Tool.')

    input_data = main.read_excel_input(input_file_location)
    table_to_print = input_data['ResourceSupply']
    table_to_print.drop_duplicates(inplace=True)

    doc.add_heading(f'Resource Supply Input Values', level=3)
    
    table = None
    for index, row in table_to_print.iloc[1:].iterrows():
        if row[0] in RESOURCE_NAME_IN_EXCEL:
            doc.add_paragraph()
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            table.autofit = True                
            hdr_cells = table.rows[0].cells
            column_names = [row[0], 'Department', 'Value', 'Informant/Data Source']
            for i, column in enumerate(column_names):
                hdr_cells[i].text = str(column)
                hdr_cells[i].paragraphs[0].runs[0].bold = True
        elif table is not None:
            row_cells = table.add_row().cells
            for i, value in enumerate(row[:4]):
                row_cells[i].text = str(value)

def add_patient_profiles_description(system, doc):

    doc.add_heading('Patient Profiles', level=2)    

    doc.add_paragraph('The following tables present the patient profiles used as the input for the MCI Planning Tool.')

    for patient_profile_name, patient_profile_info in system.components[0].patient_library.items():
        doc.add_heading(f'Patient Profile: {patient_profile_name}', level=3)
        doc.add_paragraph(f'Patient Flow: {get_patient_profile_flow(patient_profile_info)}')
        for department in patient_profile_info[:-1]:
            department_name, department_info = list(department.items())[0]
            doc.add_heading(f'{department_name}', level=4)
            doc.add_paragraph(f'Baseline Length of Stay: {department_info["BaselineLengthOfStay"]} hours')
            doc.add_paragraph(f'Baseline Mortality Rate: {department_info["BaselineMortalityRate"]*100:.2f}%')
            doc.add_paragraph('Resources Required')
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            table.autofit = True
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Resource'
            hdr_cells[0].paragraphs[0].runs[0].bold = True
            hdr_cells[1].text = 'Amount'
            hdr_cells[1].paragraphs[0].runs[0].bold = True
            hdr_cells[2].text = 'Consequence of Unmet Demand'
            hdr_cells[2].paragraphs[0].runs[0].bold = True
            hdr_cells[3].text = 'Consequence of Unmet Demand Parameter'
            hdr_cells[3].paragraphs[0].runs[0].bold = True
            for resource_info in department_info['ResourcesRequired']:
                row_cells = table.add_row().cells
                row_cells[0].text = str(resource_info['ResourceName'])
                row_cells[0].paragraphs[0].runs[0].bold = True
                if resource_info['ResourceName'] in CONSUMABLE_RESOURCES:
                    resource_amount = resource_info['ResourceAmount'] * department_info["BaselineLengthOfStay"]
                else:
                    resource_amount = resource_info['ResourceAmount']
                row_cells[1].text = str(resource_amount)
                row_cells[2].text = str(list(resource_info['ConsequencesOfUnmetDemand'][0].keys())[0])
                row_cells[3].text = str(list(resource_info['ConsequencesOfUnmetDemand'][0].values())[0])

def get_patient_profile_flow(patient_profile_info: list):
    patient_flow = ''
    for department_info in patient_profile_info[:-1]:
        patient_flow += f'{list(department_info.keys())[0]} -> '
    patient_flow += list(patient_profile_info[-1].keys())[0]
    return patient_flow    

def add_measures_of_service_summary(system, doc):

    doc.add_heading('Measures of Service Summary', level=2)

    doc.add_paragraph('The following table summarizes the measures of service for each department and patient type in the hospital.')

    departments = MCI_Planning_Tool_GUI.DEPARTMENTS
    patient_types = MCI_Planning_Tool_GUI.Ui_MainWindow.get_patient_types(system)

    for i, measure_of_service in enumerate(MEASURES_SERVICE_IN_TABLES[:3]):
        doc.add_heading(f'{MEASURES_SERVICE_IN_TABLES_LABELS[i]} in {MEASURES_SERVICE_IN_TABLES_UNITS[i]}', level=3)

        doc.add_paragraph(MEASURES_SERVICE_NOTES[measure_of_service])
        
        table = doc.add_table(rows=len(patient_types)+1, cols=len(departments)+1)
        table.style = 'Table Grid'
        table.autofit = True
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Patient Type / Department'
        for department_id, department in enumerate(departments):
            header_cells[department_id+1].text = department          
            for patient_type_id, patient_type in enumerate(patient_types):
                resilience_calculator = MCI_Planning_Tool_GUI.get_resilience_calculator(system, department.replace(' ',''), patient_type, resilience_calculator_class=ResilienceCalculator.HospitalMeasureOfServiceCalculator)
                row_cells = table.rows[patient_type_id+1].cells
                if row_cells[0].text == '':
                    patient_type_string = row_cells[0].paragraphs[0].add_run(patient_type)
                    patient_type_string.bold = True
               
                measure_of_service_string = row_cells[department_id+1].paragraphs[0].add_run(format_measure_of_service_value(resilience_calculator, measure_of_service))
                
                if measure_of_service == 'MortalityRateBefore24H' or measure_of_service == 'MortalityRateAfter24H':
                    if resilience_calculator.measures_of_service[measure_of_service] > 0:
                        measure_of_service_string.font.color.rgb = docx.shared.RGBColor(255, 0, 0)  

        for header_cell in header_cells:
            header_cell.paragraphs[0].runs[0].bold = True

    doc.add_heading('Surgeries Performed/Cancelled', level=3)

    resilience_calculator = MCI_Planning_Tool_GUI.get_resilience_calculator(system, 'All', 'All', resilience_calculator_class=ResilienceCalculator.HospitalMeasureOfServiceCalculator)

    doc.add_paragraph(f'The number of life-saving surgeries performed in the MCI during the considered period is {str(resilience_calculator.measures_of_service["SurgeriesPerformed"])}.')

    doc.add_paragraph(f'The number of life-saving surgeries cancelled in the MCI during the considered period is {str(resilience_calculator.measures_of_service["SurgeriesCancelled"])}.')
        

def format_measure_of_service_value(resilience_calculator, measure_of_service):
    if measure_of_service == 'MortalityRateBefore24H' or measure_of_service == 'MortalityRateAfter24H':
        return f'{resilience_calculator.measures_of_service[measure_of_service]*100:.2f}'
    elif measure_of_service == 'AverageLengthOfStay':
        return f'{int(resilience_calculator.measures_of_service[measure_of_service])}'
    else:
        return 'N/A'
    
def add_cause_of_death_summary(system, doc):

    doc.add_heading('Potential causes of death', level=2)

    doc.add_paragraph('The following paragraphs summarizes the potential causes of death for each patient type in the hospital by identifying their unmet resource demand. In some cases it is not clear which exact resource caused patient death as it might be that the needs of a patient for multiple resources are not met, thus all the unmet patients needs are shown here.')

    doc.add_paragraph('The unmet demand for certain resources causes death or increase in mortality rate, while others cause an increase in the length of stay. If unsure whether the unmet demand can be the potential cause of death, please check out patient profiles defined in the input file and in the Annex of the report.')

    doc.add_paragraph('Please note that insufficient supply of fuel and water results in indirect effects on patients, by decreasing the supply of oxygen, drugs and equipped beds, which then affects patients (e.g., without fuel to provide electric power, oxygen production, medical equipment and pharmacy cannot operate and provide resources to patients). However, such indirect effects are not captured explicitly. Thus, it is possible that listed cause of death, such as the lack of a bed in an OT, is not due to the insufficient number of physical beds, but due to lack of fuel (i.e., electric power) or water required to use the bed in the OT. To identify the root cause of unmet demand in such cases (i.e., lack of fuel/water or lack of physical bed), the user should consult the supply/demand plots presented below. Whether the supply of fuel/water is sufficient or might cause indirect effects can be identified in the supply/demand plots presented in the report.')

    patient_types = MCI_Planning_Tool_GUI.Ui_MainWindow.get_patient_types(system)
    patient_types.remove('All')

    resilience_calculator = MCI_Planning_Tool_GUI.get_resilience_calculator(system, 'All', 'All', resilience_calculator_class=ResilienceCalculator.CauseOfDeathCalculator)

    for patient_type in patient_types:
        for resource_name, number_of_dead_patients in resilience_calculator.unmet_demands_of_dead_patients[patient_type].items():
            line = doc.add_paragraph()
            run = line.add_run(f'{number_of_dead_patients} {patient_type} patients that died had an unmet demand for {resource_name}.')    
            run.bold = True

def add_supply_demand_plots_to_report(system, doc, department='All', patient_type='All', plot_folder='./plots/'):
    if not os.path.exists(plot_folder):
        os.makedirs(plot_folder)

    doc.add_heading('Resource Supply/Demand Plots', level=2)

    plotter_object = Plotter.Plotter()
    x_axis_label = 'Time step [hour]'
    for resource_name, resource_unit in MCI_Planning_Tool_GUI.RESOURCE_UNITS.items():
        doc.add_heading(f'Supply and Demand for {MODIFIED_RESOURCE_NAME[resource_name]}', level=3)
        y_axis_label = f'{resource_name} {resource_unit}'
        axes = plotter_object.setup_lor_plot_fig(x_axis_label, y_axis_label)
        resilience_calculator = MCI_Planning_Tool_GUI.get_resilience_calculator(system, department, patient_type, resilience_calculator_class=ResilienceCalculator.ReCoDeSResilienceCalculator)
        plotter_object.plot_single_resource(list(range(system.START_TIME_STEP, system.time_step+1)), resilience_calculator.system_supply[resource_name], 
                                                resilience_calculator.system_demand[resource_name], 
                                                resilience_calculator.system_consumption[resource_name], axes) 
        figure_name = f'{plot_folder}SupplyDemand_{resource_name.replace("/", "_")}.png'
        plotter_object.save_current_figure(figure_name, dpi=300)
        doc.add_picture(figure_name, width=docx.shared.Cm(12))        

def add_stress_scenario_description(system, doc):

    doc.add_heading('Stress Scenario', level=2)
    doc.add_paragraph(f'Stress Scenario Name: {system.damage_input.stress_scenario["StressScenarioName"]}')
    doc.add_paragraph(f'Investigated Time Period: {system.MAX_TIME_STEP+1} hours')
    generate_patient_arrival_dynamic_plots(system, doc)

def generate_patient_arrival_dynamic_plots(system, doc, plot_folder='./plots/'):
    if not os.path.exists(plot_folder):
        os.makedirs(plot_folder)
    patient_types = MCI_Planning_Tool_GUI.Ui_MainWindow.get_patient_types(system)
    patient_types.remove('All')
    existing_patients_string = f'Patients assumed to be in the hospital before the MCI are'
    existing_patients_considered = False
    for patient_type in patient_types:
        time_steps, patient_arrival = get_patient_arrival_dynamics(system, patient_type)
        if sum(patient_arrival) > 0 and time_steps != [0]:            
            plt.figure()
            plt.bar(time_steps, patient_arrival)
            plt.grid(True)
            axis_object = plt.gca()
            axis_object.yaxis.set_major_locator(matplotlib.ticker.MaxNLocator(integer=True))
            axis_object.xaxis.set_major_locator(matplotlib.ticker.MaxNLocator(integer=True))
            plt.xlabel('Time step [hour]')
            plt.ylabel('Number of patients')
            plt.title(f'Patient Arrival Dynamics for {patient_type} patients')
            plt.savefig(f'{plot_folder}PatientArrival_{patient_type.replace(" ", "_")}.png')
            doc.add_picture(f'{plot_folder}PatientArrival_{patient_type.replace(" ", "_")}.png', width=docx.shared.Cm(12))
            plt.close()

    if existing_patients_considered:        
        doc.add_paragraph(existing_patients_string[:-2] + '.')

def get_patient_arrival_dynamics(system, patient_type):
    for resource_to_change in system.damage_input.stress_scenario['ComponentsToChange'][0]['ResourcesToChange']:
        if resource_to_change['Resource'] == patient_type:
            index_of_last_non_zero = find_last_non_zero_index(resource_to_change['Amount'])
            return resource_to_change['AtTimeStep'][:index_of_last_non_zero+1], resource_to_change['Amount'][:index_of_last_non_zero+1] 

def find_last_non_zero_index(list):
    for i, element in enumerate(reversed(list)):
        if element != 0:
            return len(list) - i - 1   
    else:
        return 0

def add_liability_waiver(doc):

    doc.add_heading('Liability Waiver [To be checked with the ETH Legal Team]', level=1)

    waiver_text = [
    "The use of this software and any reports generated by it are subject to the following terms and conditions. By using this software, you agree to the following:",
    "1. **No Warranties**: The software is provided \"as is\" without any warranties, express or implied, including but not limited to the warranties of merchantability, fitness for a particular purpose, and non-infringement. The creators, developers, and distributors of this software make no representations or warranties regarding the accuracy, completeness, or reliability of the software or the reports generated by it.",
    "2. **Assumption of Risk**: You acknowledge and agree that you assume all risks associated with the use of the software and the reports generated by it. This includes, but is not limited to, any reliance on the accuracy, completeness, or usefulness of the information provided.",
    "3. **Limitation of Liability**: Under no circumstances shall the creators, developers, distributors, or any affiliates be liable for any direct, indirect, incidental, special, consequential, or exemplary damages arising from or related to the use of the software or any reports generated by it. This includes, but is not limited to, damages for loss of profits, goodwill, use, data, or other intangible losses.",
    "4. **Indemnification**: You agree to indemnify, defend, and hold harmless the creators, developers, distributors, and any affiliates from and against any claims, liabilities, damages, losses, and expenses, including without limitation, reasonable legal and accounting fees, arising out of or in any way connected with your access to or use of the software or the reports generated by it.",
    "5. **No Professional Advice**: The software and the reports generated by it are for informational purposes only and do not constitute professional advice. Users should seek professional advice before making any decisions based on the information provided by the software.",
    "6. **Governing Law**: This liability waiver shall be governed by and construed in accordance with the laws of the jurisdiction in which the software is used, without regard to its conflict of law principles.",
    "By using this software, you acknowledge that you have read, understood, and agreed to the terms of this liability waiver."
    ]   

    for line in waiver_text:
        doc.add_paragraph(line)

def add_contact_information(doc):

    doc.add_heading('Contact Information', level=1)

    doc.add_paragraph('For questions or further information, please contact the developers of the MCI Planning Tool:')

    doc.add_paragraph('Dr. Nikola Blagojevic, ETH Zurich, blagojevic@ibk.baug.ethz.ch')
    doc.add_paragraph('Prof. Dr. Bozidar Stojadinovic, ETH Zurich, stojadinovic@ibk.baug.ethz.ch')



